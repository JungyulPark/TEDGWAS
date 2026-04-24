from docx import Document

doc = Document(r'c:\ProjectTEDGWAS\RNAseqManuscript (1).docx')

# Full text
full_text = '\n'.join([p.text for p in doc.paragraphs])

# Tables
print(f"Total tables: {len(doc.tables)}")
for i, t in enumerate(doc.tables):
    header = [c.text.strip() for c in t.rows[0].cells]
    print(f"\n=== TABLE {i} ({len(t.rows)} rows x {len(t.columns)} cols) ===")
    print("Header:", header)
    for row in t.rows[1:min(20, len(t.rows))]:
        print([c.text.strip() for c in row.cells])

# Save full text for searching
with open('c:/ProjectTEDGWAS/RNAseq_fulltext.txt', 'w', encoding='utf-8') as f:
    f.write(full_text)

print("\nFull text saved.")

import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def parse_markdown_to_docx(md_path, docx_path):
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_lines = []
    
    def process_inline_formatting(paragraph, text):
        # Regular expressions for bold-italic, bold, and italic
        # We process them sequentially
        # Find all tokens: **bold**, *italic*, etc.
        pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|__.*?__|_[^_]+?_)')
        parts = pattern.split(text)
        
        for part in parts:
            if not part:
                continue
            if part.startswith('***') and part.endswith('***'):
                run = paragraph.add_run(part[3:-3])
                run.bold = True
                run.italic = True
            elif part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('__') and part.endswith('__'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('_') and part.endswith('_'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)

    def flush_table(table_lines):
        if not table_lines:
            return
        
        # Clean lines and extract rows
        rows = []
        for line in table_lines:
            # Strip outer pipes
            l = line.strip()
            if l.startswith('|'):
                l = l[1:]
            if l.endswith('|'):
                l = l[:-1]
            cells = [c.strip() for c in l.split('|')]
            rows.append(cells)
            
        # Filter out separator row (e.g., |---|---|)
        filtered_rows = []
        for r in rows:
            if all(re.match(r'^[-:]+$', c) for c in r if c):
                continue
            filtered_rows.append(r)
            
        if not filtered_rows:
            return
            
        num_cols = max(len(r) for r in filtered_rows)
        table = doc.add_table(rows=len(filtered_rows), cols=num_cols)
        table.style = 'Table Grid'
        
        # Populate table
        for r_idx, r in enumerate(filtered_rows):
            row_cells = table.rows[r_idx].cells
            for c_idx, val in enumerate(r):
                if c_idx < len(row_cells):
                    # Process inline formatting for cell
                    cell_p = row_cells[c_idx].paragraphs[0]
                    process_inline_formatting(cell_p, val)
                    if r_idx == 0:
                        cell_p.runs[0].font.bold = True if cell_p.runs else False
                        
        doc.add_paragraph() # Add spacing after table

    for idx, line in enumerate(lines):
        stripped = line.strip()
        
        # Check if we are in a table
        if stripped.startswith('|'):
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
            continue
        else:
            if in_table:
                flush_table(table_lines)
                in_table = False
                table_lines = []
        
        # Empty lines
        if not stripped:
            continue
            
        # Horizontal rules
        if stripped == '---':
            doc.add_paragraph().add_run('━' * 50).font.color.rgb = None
            continue
            
        # Headers
        if stripped.startswith('# '):
            heading = doc.add_paragraph()
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)
            run = heading.add_run(stripped[2:])
            run.font.size = Pt(16)
            run.bold = True
            continue
        elif stripped.startswith('## '):
            heading = doc.add_paragraph()
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)
            run = heading.add_run(stripped[3:])
            run.font.size = Pt(14)
            run.bold = True
            continue
        elif stripped.startswith('### '):
            heading = doc.add_paragraph()
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)
            run = heading.add_run(stripped[4:])
            run.font.size = Pt(12)
            run.bold = True
            continue
            
        # Lists (unordered or ordered)
        list_match = re.match(r'^([\*\-\+])\s+(.*)$', stripped)
        num_list_match = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        
        if list_match:
            p = doc.add_paragraph(style='List Bullet')
            process_inline_formatting(p, list_match.group(2))
        elif num_list_match:
            p = doc.add_paragraph(style='List Number')
            process_inline_formatting(p, num_list_match.group(2))
        else:
            # Paragraph
            p = doc.add_paragraph()
            # If the paragraph is blockquote (starts with >)
            if stripped.startswith('>'):
                p.paragraph_format.left_indent = Inches(0.5)
                stripped_text = stripped[1:].strip()
                process_inline_formatting(p, stripped_text)
            else:
                process_inline_formatting(p, stripped)
                
    # If file ends while in table
    if in_table:
        flush_table(table_lines)
        
    doc.save(docx_path)
    print(f"Successfully compiled {md_path} to {docx_path}")

if __name__ == '__main__':
    parse_markdown_to_docx(
        r'c:\ProjectTEDGWAS\MANUSCRIPT_TED_TRAP_SUBMISSION.md',
        r'c:\ProjectTEDGWAS\MANUSCRIPT_TED_TRAP_SUBMISSION.docx'
    )

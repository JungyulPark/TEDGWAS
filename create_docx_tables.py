import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Read frozen data
excel_path = r'c:\ProjectTEDGWAS\Final_Numbers_Frozen_20260422_v2.xlsx'
study_design = pd.read_excel(excel_path, sheet_name='Study_Design')
mr_results = pd.read_excel(excel_path, sheet_name='MR_Results')
mr_sens = pd.read_excel(excel_path, sheet_name='MR_Sensitivity')

doc = Document()
doc.add_heading('TED-GWAS Manuscript Tables (Phase B)', 0)

# Set up some styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

def set_cell_border(cell, **kwargs):
    """
    Set cell's border
    Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "#000000", "space": "0"}, ...)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

# ================= Table 1: Study Design =================
doc.add_heading('Table 1. Characteristics of GWAS data and cohorts used in the study.', level=1)
t1_cols = ['Category', 'GWAS ID', 'Trait', 'Total N', 'Cases', 'Controls', 'Population', 'Year']
t1 = doc.add_table(rows=1, cols=len(t1_cols))
t1.style = 'Table Grid'

hdr_cells = t1.rows[0].cells
for i, col_name in enumerate(t1_cols):
    hdr_cells[i].text = col_name
    hdr_cells[i].paragraphs[0].runs[0].font.bold = True

for _, r in study_design.iterrows():
    row_cells = t1.add_row().cells
    row_cells[0].text = str(r['Category'])
    row_cells[1].text = str(r['GWAS_ID'])
    row_cells[2].text = str(r['Trait'])
    row_cells[3].text = str((int(r['N_Total']) if pd.notna(r['N_Total']) and r['N_Total'] != '' else '-'))
    row_cells[4].text = str((int(r['N_Cases']) if pd.notna(r['N_Cases']) and r['N_Cases'] != '' else '-'))
    row_cells[5].text = str((int(r['N_Controls']) if pd.notna(r['N_Controls']) and r['N_Controls'] != '' else '-'))
    row_cells[6].text = str(r['Population'])
    row_cells[7].text = str(int(r['Year']) if pd.notna(r['Year']) and str(r['Year']).isdigit() else "-")
    
doc.add_paragraph()

# ================= Table 2: MR Results =================
doc.add_heading('Table 2. Mendelian Randomization (MR) estimates of genetically predicted gene expression on TED and associated traits.', level=1)
# Join MR results
mr_all = mr_results.copy()

t2_cols = ['Gene', 'Outcome', 'β (SE)', 'P-value', 'nSNP', 'Sensitivity Notes']
t2 = doc.add_table(rows=1, cols=len(t2_cols))
t2.style = 'Table Grid'

hdr_cells2 = t2.rows[0].cells
for i, col_name in enumerate(t2_cols):
    hdr_cells2[i].text = col_name
    hdr_cells2[i].paragraphs[0].runs[0].font.bold = True

# Format rules: group by Gene
merged_mr = pd.merge(mr_results, mr_sens, on=['Gene', 'Outcome'], how='left')

for gene in ['TSHR', 'IGF1R', 'TNF', 'PPARG', 'ARRB1', 'CTLA4', 'IRS1', 'AKT1']:
    gene_data = merged_mr[merged_mr['Gene'] == gene]
    if len(gene_data) == 0: continue
    
    for idx, (_, r) in enumerate(gene_data.iterrows()):
        row_cells = t2.add_row().cells
        
        # Only print gene name on first row of group
        row_cells[0].text = gene if idx == 0 else ""
        if idx == 0:
            row_cells[0].paragraphs[0].runs[0].font.bold = True
        
        row_cells[1].text = str(r['Outcome']).replace('Primary(Graves)', 'Primary').replace('Replication(Hyper)', 'Replication').replace('FinnGen_Sensitivity', 'Sensitivity (FinnGen)')
        beta = float(r['beta'])
        se = float(r['SE'])
        
        row_cells[2].text = f"{beta:.3f} ({se:.3f})"
        
        pval_val = float(r['P_value'])
        if pval_val < 0.001:
            pval_str = f"{pval_val:.2e}"
        else:
            pval_str = f"{pval_val:.3f}"
        
        row_cells[3].text = pval_str
        if pval_val < 0.05:
            row_cells[3].paragraphs[0].runs[0].font.bold = True
            
        row_cells[4].text = str(r['nSNP_x'])
        
        notes = str(r['Notes']) if pd.notna(r['Notes']) else ""
        if r['Egger_Intercept_P'] and pd.notna(r['Egger_Intercept_P']) and str(r['Egger_Intercept_P']) != 'N/A (nSNP<3)':
            notes += f" Egger P={float(r['Egger_Intercept_P']):.3f}"
        row_cells[5].text = notes.strip()

out_path = r'c:\ProjectTEDGWAS\TED_TRAP_Tables_1_2_v1.docx'
doc.save(out_path)
print(f"Tables successfully written to {out_path}")

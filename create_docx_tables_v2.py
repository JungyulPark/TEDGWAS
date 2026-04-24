import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Read frozen data
excel_path = r'c:\ProjectTEDGWAS\Final_Numbers_Frozen_20260422_v2.xlsx'
study_design = pd.read_excel(excel_path, sheet_name='Study_Design')
mr_results = pd.read_excel(excel_path, sheet_name='MR_Results')
mr_sens = pd.read_excel(excel_path, sheet_name='MR_Sensitivity')

doc = Document()
doc.add_heading('TED-GWAS Manuscript Tables (Phase B)', 0)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

def format_number(val):
    if pd.isna(val): return ""
    v = float(val)
    if v == 0: return "0.000"
    if abs(v) < 0.001:
        return f"{v:.2e}"
    else:
        return f"{v:.3f}"

# ================= Table 1: Study Design =================
doc.add_heading('Table 1. Characteristics of GWAS data and cohorts used in the study.', level=1)
t1_cols = ['Category', 'GWAS ID', 'Trait', 'Total N', 'Cases', 'Controls', 'Population', 'Year']
t1 = doc.add_table(rows=1, cols=len(t1_cols))
t1.style = 'Table Grid'

hdr_cells = t1.rows[0].cells
for i, col_name in enumerate(t1_cols):
    hdr_cells[i].text = col_name
    hdr_cells[i].paragraphs[0].runs[0].font.bold = True

for _, r in study_design.iterrows():
    row_cells = t1.add_row().cells
    row_cells[0].text = str(r['Category'])
    row_cells[1].text = str(r['GWAS_ID'])
    row_cells[2].text = str(r['Trait'])
    
    val_total = r['N_Total']
    row_cells[3].text = str(int(val_total)) if pd.notna(val_total) and str(val_total).strip() != '' else '-'
    
    val_cases = r['N_Cases']
    row_cells[4].text = str(int(val_cases)) if pd.notna(val_cases) and str(val_cases).strip() != '' else '-'
    
    val_ctrls = r['N_Controls']
    row_cells[5].text = str(int(val_ctrls)) if pd.notna(val_ctrls) and str(val_ctrls).strip() != '' else '-'
    
    row_cells[6].text = str(r['Population'])
    
    if r['Category'] == 'RNA-seq':
        row_cells[7].text = 'Unpublished'
    else:
        year_val = r['Year']
        row_cells[7].text = str(int(year_val)) if pd.notna(year_val) and str(year_val).replace('.','',1).isdigit() else '-'

doc.add_paragraph()

# ================= Table 2: MR Results =================
doc.add_heading('Table 2. Mendelian Randomization (MR) estimates of genetically predicted gene expression on TED and associated traits.', level=1)
t2_cols = ['Gene', 'Outcome', 'β (SE)', 'P-value', 'nSNP', 'Sensitivity Notes']
t2 = doc.add_table(rows=1, cols=len(t2_cols))
t2.style = 'Table Grid'

# Set column widths (docx layout)
widths = [Inches(1.0), Inches(2.0), Inches(1.5), Inches(1.0), Inches(0.8), Inches(2.0)]
for row in t2.rows:
    for idx, width in enumerate(widths):
        row.cells[idx].width = width

hdr_cells2 = t2.rows[0].cells
for i, col_name in enumerate(t2_cols):
    hdr_cells2[i].text = col_name
    hdr_cells2[i].paragraphs[0].runs[0].font.bold = True
    hdr_cells2[i].width = widths[i]

merged_mr = pd.merge(mr_results, mr_sens, on=['Gene', 'Outcome'], how='left')

for gene in ['TSHR', 'IGF1R', 'TNF', 'PPARG', 'ARRB1', 'CTLA4', 'IRS1', 'AKT1']:
    gene_data = merged_mr[merged_mr['Gene'] == gene]
    if len(gene_data) == 0: continue
    
    for idx, (_, r) in enumerate(gene_data.iterrows()):
        row_cells = t2.add_row().cells
        
        # widths
        for i, w in enumerate(widths):
            row_cells[i].width = w
            
        row_cells[0].text = gene if idx == 0 else ""
        if idx == 0:
            row_cells[0].paragraphs[0].runs[0].font.bold = True
        
        outcome_str = str(r['Outcome'])
        outcome_str = outcome_str.replace('Primary(Graves)', 'Primary').replace('Replication(Hyper)', 'Replication').replace('FinnGen_Sensitivity', 'Sensitivity (FinnGen)')
        row_cells[1].text = outcome_str
        
        beta = float(r['beta'])
        se = float(r['SE'])
        
        beta_str = format_number(beta)
        se_str = format_number(se)
        
        row_cells[2].text = f"{beta_str} ({se_str})"
        
        pval_val = float(r['P_value'])
        row_cells[3].text = format_number(pval_val)
        if pval_val < 0.05:
            row_cells[3].paragraphs[0].runs[0].font.bold = True
            
        row_cells[4].text = str(r['nSNP_x'])
        
        notes = str(r['Notes']) if pd.notna(r['Notes']) else ""
        if 'Egger_Intercept_P' in r and pd.notna(r['Egger_Intercept_P']):
            ep = str(r['Egger_Intercept_P'])
            if 'N/A' not in ep:
                notes += f" Egger P={float(ep):.3f}"
        row_cells[5].text = notes.strip()

out_path = r'c:\ProjectTEDGWAS\TED_TRAP_Tables_1_2_v1.docx'
doc.save(out_path)
print(f"Tables successfully fixed and written to {out_path}")

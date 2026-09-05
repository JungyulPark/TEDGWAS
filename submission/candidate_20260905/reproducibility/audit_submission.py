import sys,re,json,hashlib
from pathlib import Path
W=Path(__file__).resolve().parent;sys.path.insert(0,str(W/'pydeps'))
import numpy as np,pandas as pd
from scipy.stats import norm
from docx import Document
O=W.parent;P=O/'provenance'
t=(O.parents[1]/'MANUSCRIPT_TED_TRAP_v5_MASTER.md').read_text(encoding='utf-8')
co=pd.read_csv(P/'coloc_canonical_v2.csv');mr=pd.read_csv(P/'MR_primary_canonical.csv');inst=pd.read_csv(P/'instruments_verified.csv')
raw=pd.read_csv(P/'MR_all_estimators_verified.csv')
errors=[];checks=[];n=0
def check(ok,label):
    checks.append(label)
    if not ok:errors.append(label)
def cells(line):return [x.strip().replace('*','') for x in re.split(r'(?<!\\)\|',line.strip().strip('|'))]
tables={};current=None;rows=[];serial=0
for line in t.splitlines()+['']:
    match=re.match(r'(?:- )?\*\*Table (S?\d+)\.',line)
    if match:current=match[1]
    if line.startswith('|'):
        if re.match(r'^\|[-:| ]+\|$',line):continue
        rows.append(cells(line))
    elif rows:
        key=current
        if key in tables:key=key+'b'
        tables[key]=rows;rows=[]
ocmap={'BBJ Graves':'BBJ_Graves','UKB hyperthyroid':'UKB_hyperthyroid','FinnGen GO':'FinnGen_GO'}
trans=str.maketrans('⁰¹²³⁴⁵⁶⁷⁸⁹⁻−','0123456789--')
def numeric(s):return float(s.replace(',','').replace('×10','e').translate(trans))
def agree(shown,actual,label):
    global n
    n+=1;s=shown.replace('*','').strip()
    if s.startswith('<'):
        check(actual<numeric(s[1:]),label);return
    v=numeric(s)
    if '×10' in s:
        base,exp=s.translate(trans).split('×10');unit=10**int(exp)*10**(-len(base.split('.')[1]) if '.' in base else 0)
    else:unit=10**(-len(s.split('.')[1]) if '.' in s else 0)
    check(abs(v-actual)<=unit/2+1e-12,f'{label}: displayed={s}, source={actual:.12g}')
def getco(g,oc,p=1e-5):return co[(co.gene==g)&(co.outcome==oc)&np.isclose(co.p12,p,atol=1e-12)].iloc[0]
def getmr(g,oc):return mr[(mr.gene_symbol==g)&(mr.outcome==oc)].iloc[0]
for key in ['S2','S9']:
    for row in tables[key][1:]:
        g,oc=row[0],ocmap[row[1]];r=getco(g,oc)
        check(int(row[2].replace(',',''))==int(r.n_overlap),f'{key} {g}/{oc} overlaps')
        check(row[8]==r.top_snp,f'{key} {g}/{oc} lead')
        for i in range(5):agree(row[3+i],r[f'PP.H{i}'],f'{key} {g}/{oc} H{i}')
        for j,p in [(9,1e-6),(10,5e-6)]:agree(row[j],getco(g,oc,p)['PP.H4'],f'{key} {g}/{oc} p12={p}')
g=None
for row in tables['2'][1:]:
    g=row[0] or g;oc=ocmap[row[1]];r=getco(g,oc);m=getmr(g,oc)
    check(int(row[2])==int(m.n_iv),f'Table2 {g}/{oc} IVs')
    agree(row[4],m.pvalue,f'Table2 {g}/{oc} P')
    for j,i in [(8,4),(9,3)]:agree(row[j],r[f'PP.H{i}'],f'Table2 {g}/{oc} H{i}')
    check(row[10]==r.top_snp,f'Table2 {g}/{oc} lead')
    nums=re.findall(r'[0-9]+\.[0-9]+',row[3]);
    for v,col in zip(nums,['or','or_ci_lower','or_ci_upper']):agree(v,m[col],f'Table2 {g}/{oc} {col}')
for row in tables['3'][1:]:
    g=row[0];m=getmr(g,'BBJ_Graves');agree(row[3],m.pvalue,f'Table3 {g} P')
    check(int(row[1])==int(m.n_iv),f'Table3 {g} IVs')
    for v,col in zip(re.findall(r'[0-9]+\.[0-9]+',row[2]),['or','or_ci_lower','or_ci_upper']):agree(v,m[col],f'Table3 {g} {col}')
    for j,oc in enumerate(ocmap.values(),4):
        if row[j]!='—':agree(row[j],getco(g,oc)['PP.H4'],f'Table3 {g}/{oc} H4')
for row in tables['S5'][1:]:
    for j,oc in enumerate(ocmap.values(),4):agree(row[j],getco(row[0],oc)['PP.H4'],f'S5 {row[0]}/{oc} H4')
for row in tables['S1'][1:]:
    for j,oc in enumerate(ocmap.values(),1):
        if row[j]=='—':
            check(mr[(mr.gene_symbol==row[0])&(mr.outcome==oc)].empty,f'S1 {row[0]}/{oc} absent estimate');continue
        pair=re.match(r'(.+?)\s*\((.+)\)',row[j]);m=getmr(row[0],oc)
        agree(pair[1],m.beta,f'S1 {row[0]}/{oc} beta');agree(pair[2],m.pvalue,f'S1 {row[0]}/{oc} P')
for row in tables['S4'][1:]:check(int(row[2])==len(inst[inst.gene_symbol==row[0]]),f'S4 {row[0]} instrument count')
for row in tables['S6'][1:]:
    g,oc=row[0],ocmap[row[1]];m=getmr(g,oc);agree(row[3],m.pvalue,f'S6 {g}/{oc} P')
    for j,method in [(4,'Weighted median'),(5,'Weighted mode')]:
        if row[j]!='NA':agree(row[j],raw[(raw.gene_symbol==g)&(raw.outcome==oc)&(raw.method==method)].iloc[0].pvalue,f'S6 {g}/{oc} {method}')
    for j,col in [(6,'egger_intercept_p'),(7,'cochran_q_p')]:
        if row[j]!='NA':agree(row[j],m[col],f'S6 {g}/{oc} {col}')
g=None
for row in tables['S8b'][1:]:
    g=row[0] or g;oc=ocmap[row[1]];m=getmr(g,oc);alpha=.05/2544 if oc=='BBJ_Graves' else .05;v=(norm.ppf(1-alpha/2)+norm.ppf(.8))*m.se
    for j,val in [(2,m.beta),(3,m.se),(4,v),(5,np.exp(v))]:agree(row[j],val,f'S8b {g}/{oc} column{j}')
check(len(inst)==6135 and inst.gene_symbol.nunique()==2544,'Verified instrument denominator 6135/2544')
check(dict(mr.groupby('outcome').size())=={'BBJ_Graves':2234,'FinnGen_GO':2480,'UKB_hyperthyroid':2505},'Outcome counts after GAN exclusion')
check(len(mr[(mr.outcome=='BBJ_Graves')&(mr.pvalue<.05/2544)])==13,'13 discovery hits')
check(set(tables)=={'1','2','3','S1','S2','S3','S4','S5','S6','S7','S8','S8b','S9'},'All 13 physical tables present in Markdown')
check(np.max(abs(co[[f'PP.H{i}' for i in range(5)]].sum(axis=1)-1))<1e-12,'Posterior simplex for 81 rows')
ver=json.loads((P/'coloc_verification.json').read_text());check(ver['canonical_sha256']==hashlib.sha256((P/'coloc_canonical_v2.csv').read_bytes()).hexdigest(),'Independent reproduction bound to current canonical CSV hash')
check(ver['max_abs_posterior_difference']<1e-9 and ver['rows']==81,'Independent 81-row posterior reproduction')
check('sensitivity analysis using eQTLGen\'s own allele frequencies has not been performed' in t,'Unperformed frequency sensitivity remains disclosed')
for name,cnt in [('MANUSCRIPT_Submission',3),('SUPPLEMENTARY_MATERIAL',10)]:
    doc=Document(O/(name+'.docx'));check(len(doc.tables)==cnt,name+' table count')
    drows=[[c.text for c in row.cells] for tb in doc.tables for row in tb.rows]
    check(all('�' not in c for row in drows for c in row),name+' no replacement glyphs')
    check(all(tb.rows[0]._tr.xpath('./w:trPr/w:tblHeader') for tb in doc.tables),name+' repeating headers')
    check(all(row._tr.xpath('./w:trPr/w:cantSplit') for tb in doc.tables for row in tb.rows),name+' unsplit table rows')
check('0.799' in '\n'.join(c for row in [[c.text for c in r.cells] for tb in Document(O/'MANUSCRIPT_Submission.docx').tables for r in tb.rows] for c in row),'DOCX includes final CTLA4 H3')
res={'status':'PASS' if not errors else 'FAIL','numeric_cells_compared':n,'checks':len(checks),'errors':errors,'scope':'Source-to-table consistency and DOCX structural checks; author declarations and unperformed sensitivity are not validated by this pass.'}
(P/'integrity_audit.json').write_text(json.dumps(res,indent=2),encoding='utf-8')
print(json.dumps(res,indent=2));sys.exit(bool(errors))

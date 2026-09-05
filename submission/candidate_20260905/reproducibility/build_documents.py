import sys,re,subprocess,copy,shutil,json
from pathlib import Path
W=Path(__file__).resolve().parent;sys.path.insert(0,str(W/'pydeps'))
from docx import Document
from docx.shared import Inches,Pt,RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START,WD_ORIENT
O=W.parent;P=O/'provenance'
master=O.parents[1]/'MANUSCRIPT_TED_TRAP_v5_MASTER.md';t=master.read_text(encoding='utf-8')
t=t.replace('**Figure 2. Discovery Mendelian randomization across the druggable genome.**','**Figure 2. Discovery Mendelian randomization across the druggable genome.** **(A)** All estimable genes. **(B)** Enlarged effect range containing all thirteen discovery hits.')
extra='Points represent individual samples; horizontal lines mark the mean TED abundance. Log2 fold changes use a 0.01-TPM pseudocount. '
t=t.replace(extra,'')
t=t.replace('**Figure S1. Orbital tissue transcript abundance (descriptive).**','**Figure S1. Orbital tissue transcript abundance (descriptive).** '+extra.rstrip())
t=t.replace('This framework supports *TSHR* as an expression-colocalized susceptibility anchor for GD, whereas the *IGF1R* profile is more compatible with pharmacologic effector biology than with the *TSHR*-like expression-colocalized susceptibility architecture, while not excluding a more modest genetic contribution.', 'These results support *TSHR* as an expression-colocalized susceptibility anchor for GD, while shared-variant evidence at *IGF1R* remains unresolved. The differing genetic profiles do not establish an exclusive effector role for *IGF1R* or exclude a modest inherited contribution.')
t=t.replace('the posterior never approached the shared-variant threshold','the posterior remained below the shared-variant threshold')
t=t.replace('**Figure 2. Discovery Mendelian randomization across the druggable genome.** **(A)** All estimable genes. **(B)** Enlarged effect range containing all thirteen discovery hits. **(A)** All estimable genes. **(B)** Enlarged effect range containing all thirteen discovery hits.','**Figure 2. Discovery Mendelian randomization across the druggable genome.** **(A)** All estimable genes. **(B)** Enlarged effect range containing all thirteen discovery hits.')
# An ASCII absolute-value label avoids pipe-table ambiguity.
t=t.replace('\\|β\\|min (80% power)','Absolute βmin (80% power)')
t=t.replace('| *IFNGR1* | 6 | 1 | 340 | Yes |','| *IFNGR1* | 6 | 2 | 340 | No |')
t=t.replace('| +0.14 (0.68) | rs11754268 |','| +0.15 (0.68) | rs11754268 |')
t=t.replace('| 0.021 | 0.029 | 0.118 |','| 0.021 | 0.029 | 0.512 |').replace('| 0.012 | 0.002 | 0.040 |','| 0.012 | 0.002 | 0.367 |').replace('| 0.182 | 0.124 | 0.241 |','| 0.182 | 0.124 | 0.633 |')
# Avoid nested asterisk delimiters in long table footnotes and isolate every caption.
t='\n'.join(line[:len(line)-len(line.lstrip())]+line.strip()[1:-1] if re.match(r'^\s{0,2}\*[^*].*\*\s*$',line) else line for line in t.split('\n'))
t=re.sub(r'\n(?:- )?(\*\*Table S)',r'\n\n\1',t)
master.write_text(t,encoding='utf-8')
main,sup=t.split('## Supplementary Material',1)
main=main.rstrip('-\n ')+'\n'
title=t.split('\n# ')[1].split('\n')[0]
sup='# Supplementary material\n\n'+title+'\n\n'+sup.strip()+'\n'
sup=re.sub(r'(?m)^- (\*\*Table S\d)',r'\1',sup)
sup=re.sub(r'(?m)^  (?! )','',sup)
cover=f'''# Cover letter

5 September 2026

The Editors  
*Endocrine Connections*

Dear Editors,

Please consider our manuscript, “{title},” as an original research article.

The efficacy of IGF-1R blockade in thyroid eye disease raises a clinically relevant question: how does the genetic susceptibility profile of *IGF1R* compare with that of the established Graves disease autoantigen *TSHR*? We examined genetically proxied blood expression across 2,544 druggable genes using Mendelian randomization and colocalization. Biobank Japan Graves disease served as discovery, UK Biobank hyperthyroidism as broader-phenotype replication, and FinnGen Graves ophthalmopathy as a TED-enriched sensitivity outcome.

Thirteen genes reached discovery significance. *TSHR* showed direction-consistent protective MR estimates and strong expression colocalization in BBJ and FinnGen (PP.H4 = 0.951 and 0.986), while UKB favoured distinct variants. *IGF1R* had nominal MR associations in BBJ and UKB but did not reach the colocalization threshold in any outcome. No novel candidate passed the combined discovery-plus-TED colocalization criterion. The study offers a common framework for interpreting these profiles without equating therapeutic efficacy with genetic susceptibility or inferring that an unresolved locus has no genetic contribution.

The manuscript reports limitations that affect clinical interpretation: FinnGen uses population controls and cannot isolate TED-specific liability within Graves disease; the exposure data are European blood eQTLs; colocalization depends on the single-causal-variant model and prior; and cohort-specific eQTLGen frequency sensitivity remains untested. MR detection thresholds are reported separately from the unquantified power of the combined filter. Orbital RNA-seq observations from four TED samples and one control are descriptive only, with no inferential expression tests reported.

An earlier, narrower version was considered at thyroid-focused journals. Following editorial feedback, the work was restructured around a druggable-gene-wide discovery analysis and multi-outcome evaluation. The in-house orbital dataset is also used for a different, whole-transcriptome question in a separate manuscript under review elsewhere. Here, only the three backbone genes are described; no genome-wide expression catalogue or pathway-enrichment result is presented. The companion manuscript can be provided to the editors to assess overlap.

This manuscript has not been published and is not under consideration elsewhere. The authors report no specific funding or competing interests. The orbital tissue component received approval from the Institutional Review Board of Pusan National University Hospital (2104-018-102), with written informed consent. Author contributions, data access, reporting coverage and use of AI-assisted tools are disclosed in the manuscript. Instrument-level data and primary MR and colocalization results accompany the submission.

Thank you for considering this work.

Sincerely,

Suk-Woo Yang, MD, PhD  
Corresponding author, on behalf of Jungyul Park, Min-Seon Kim, Kyung-Hwa Shin and Suk-Woo Yang  
Department of Ophthalmology, Seoul St. Mary's Hospital  
College of Medicine, The Catholic University of Korea, Seoul, Republic of Korea  
E-mail: yswoph@catholic.ac.kr · Tel: +82-2-2258-2847
'''
cover_path=O/'COVER_LETTER_EndocrineConnections.md'
if cover_path.exists():cover=cover_path.read_text(encoding='utf-8')
cover_path.write_text(cover,encoding='utf-8')
texts={'MANUSCRIPT_Submission':main,'SUPPLEMENTARY_MATERIAL':sup,'COVER_LETTER_EndocrineConnections':cover}
pandoc=shutil.which('pandoc')
if not pandoc:
    import pypandoc
    pandoc=pypandoc.get_pandoc_path()
reference=W/'submission_reference.docx';Document().save(reference)
def setsec(sec,land=False,lines=False):
    sec.orientation=WD_ORIENT.LANDSCAPE if land else WD_ORIENT.PORTRAIT
    sec.page_width=Inches(11.69 if land else 8.27);sec.page_height=Inches(8.27 if land else 11.69)
    sec.left_margin=sec.right_margin=Inches(.7 if land else .95)
    sec.top_margin=sec.bottom_margin=Inches(.65 if land else .8)
    sec.header_distance=sec.footer_distance=Inches(.3)
    for old in sec._sectPr.findall(qn('w:lnNumType')):sec._sectPr.remove(old)
    if lines:
        x=OxmlElement('w:lnNumType');x.set(qn('w:countBy'),'1');x.set(qn('w:restart'),'continuous');x.set(qn('w:distance'),'240');sec._sectPr.append(x)
def boundary_before(doc,p,land_previous):
    # Section properties on a paragraph describe the section ending there.
    new=OxmlElement('w:p');pr=OxmlElement('w:pPr');sp=copy.deepcopy(doc.sections[-1]._sectPr)
    for x in sp.findall(qn('w:type')):sp.remove(x)
    typ=OxmlElement('w:type');typ.set(qn('w:val'),'nextPage');sp.append(typ)
    # Geometry for the ending section.
    sz=sp.find(qn('w:pgSz'));sz.set(qn('w:w'),str(int((11.69 if land_previous else 8.27)*1440)));sz.set(qn('w:h'),str(int((8.27 if land_previous else 11.69)*1440)))
    if land_previous:sz.set(qn('w:orient'),'landscape')
    else:sz.attrib.pop(qn('w:orient'),None)
    mar=sp.find(qn('w:pgMar'))
    for side in ['left','right']:mar.set(qn('w:'+side),str(int((.7 if land_previous else .95)*1440)))
    for side in ['top','bottom']:mar.set(qn('w:'+side),str(int((.65 if land_previous else .8)*1440)))
    for x in sp.findall(qn('w:lnNumType')):sp.remove(x)
    if not land_previous and 'MANUSCRIPT' in doc.core_properties.identifier:
        ln=OxmlElement('w:lnNumType');ln.set(qn('w:countBy'),'1');ln.set(qn('w:restart'),'continuous');ln.set(qn('w:distance'),'240');sp.append(ln)
    pr.append(sp);new.append(pr);p._p.addprevious(new)
for name,txt in texts.items():
    # Remove horizontal-rule paragraphs and make source supplementary captions normal paragraphs.
    txt=re.sub(r'(?m)^---\s*$','',txt)
    md=W/(name+'.md');md.write_text(txt,encoding='utf-8')
    out=O/(name+'.docx')
    subprocess.run([str(pandoc),str(md),'-f','markdown+pipe_tables+tex_math_dollars','-t','docx','-o',str(out)],check=True)
    d=Document(out);d.core_properties.title=title if name!='COVER_LETTER_EndocrineConnections' else 'Cover letter to Endocrine Connections';d.core_properties.identifier=name;d.core_properties.author='Jungyul Park; Min-Seon Kim; Kyung-Hwa Shin; Suk-Woo Yang'
    ismain=name=='MANUSCRIPT_Submission';iscover=name.startswith('COVER');issup=name=='SUPPLEMENTARY_MATERIAL'
    for sec in d.sections:setsec(sec,False,ismain)
    for stname in ['Normal','Body Text','First Paragraph','Compact','Caption','Table','List Paragraph','Footnote Text']:
        if stname not in d.styles:continue
        st=d.styles[stname];st.font.name='Times New Roman';st.font.size=Pt(11 if iscover else 12)
        st.paragraph_format.line_spacing=1.0 if iscover else 2
        st.paragraph_format.space_after=Pt(5 if iscover else 0)
    for nm in ['Title','Heading 1','Heading 2','Heading 3','Heading 4']:
        st=d.styles[nm];st.font.name='Times New Roman';st.font.color.rgb=RGBColor(0,0,0);st.font.bold=True;st.font.size=Pt(15 if nm=='Title' else 12)
        st.paragraph_format.space_before=Pt(12);st.paragraph_format.space_after=Pt(6);st.paragraph_format.keep_with_next=True
        st.paragraph_format.line_spacing=1.15
    for p in d.paragraphs:
        p.paragraph_format.widow_control=True
        if p.style.name=='Heading 1' and (p.text==title or p.text=='Supplementary material' or p.text=='Cover letter'):p.style='Title'
        if p.text in ['Abstract','Introduction','References','Figure Legends']:p.paragraph_format.page_break_before=True
        if p.text.startswith('Table ') or p.text.startswith('Panel '):
            p.paragraph_format.keep_with_next=True;p.paragraph_format.line_spacing=1.1;p.paragraph_format.space_after=Pt(7)
        if p.text.startswith('Table S') or re.match(r'^Table [123]\.',p.text):
            if (ismain and p.text.startswith('Table 1.')) or (issup and p.text.startswith('Table S1.')):
                boundary_before(d,p,False)
            else:p.paragraph_format.page_break_before=True
        if p.text=='Tables':p._element.getparent().remove(p._element)
        if iscover:
            p.paragraph_format.widow_control=True
            if p.text=='Cover letter':p.paragraph_format.space_before=Pt(0);p.paragraph_format.space_after=Pt(6)
            if p.text.startswith('Sincerely') or p.text.startswith('Suk-Woo Yang'):p.paragraph_format.keep_with_next=True
        # Single-spaced reference entries improve usability while body remains double spaced.
        if p.style.name.startswith('List') or re.match(r'^\d+\. ',p.text):
            p.paragraph_format.line_spacing=1;p.paragraph_format.space_after=Pt(6)
    if not iscover:
        setsec(d.sections[-1],True,False)
        # Landscape table footnotes are compact but retain source text and italics.
        in_tables=False
        for p in d.paragraphs:
            if p.text.startswith('Table 1.') or p.text.startswith('Table S1.'):in_tables=True
            if in_tables:
                p.paragraph_format.line_spacing=1.05;p.paragraph_format.space_after=Pt(6)
                for r in p.runs:r.font.size=Pt(10)
    # Reconstruct tables without Pandoc table-style and width inheritance.
    for old in list(d.tables):
        fresh=d.add_table(rows=len(old.rows),cols=len(old.columns))
        for source_row,target_row in zip(old.rows,fresh.rows):
            for source_cell,target_cell in zip(source_row.cells,target_row.cells):
                for el in list(target_cell._tc):
                    if el.tag!=qn('w:tcPr'):target_cell._tc.remove(el)
                for sp in source_cell.paragraphs:target_cell._tc.append(copy.deepcopy(sp._p))
        old._tbl.addprevious(fresh._tbl);old._tbl.getparent().remove(old._tbl)
    for table in d.tables:
        table.autofit=False;table.alignment=WD_TABLE_ALIGNMENT.CENTER
        cols=len(table.columns);available=10.29
        if cols==11:weights=[.72,1.1,.78,.62,.62,.62,.62,.62,1.0,.93,.93] if 'Overlapping' in table.cell(0,2).text else [.6,1.04,.32,1.05,.75,.76,.92,.82,.6,.6,1.0]
        elif cols==3:weights=[.45,1.65,8.19]
        elif cols==8 and 'Classification' in table.cell(0,7).text:weights=[.72,.33,1.1,.74,.8,.8,.9,2.6]
        elif cols==8 and table.cell(0,0).text=='Dataset':weights=[1.5,1.3,1.45,1.1,.65,.8,.8,.7]
        elif cols==6:weights=[.8,.4,1.6,1.65,.8,1.1] if table.cell(0,1).text=='Chr' else [.7,1.15,1.15,1.15,1.05,2.2]
        else:weights=[1]*cols
        widths=[available*x/sum(weights) for x in weights]
        tw=table._tbl.tblPr.find(qn('w:tblW'));tw.set(qn('w:type'),'dxa');tw.set(qn('w:w'),str(round(available*1440)))
        for col,width in zip(table.columns,widths):col.width=Inches(width)
        for ri,row in enumerate(table.rows):
            trpr=row._tr.get_or_add_trPr();cant=OxmlElement('w:cantSplit');trpr.append(cant)
            if ri==0:repeat=OxmlElement('w:tblHeader');trpr.append(repeat)
            for ci,cell in enumerate(row.cells):
                cell.width=Inches(widths[ci]);cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
                pr=cell._tc.get_or_add_tcPr();mar=OxmlElement('w:tcMar')
                for side,val in [('top','65'),('bottom','65'),('left','65'),('right','65')]:x=OxmlElement('w:'+side);x.set(qn('w:w'),val);x.set(qn('w:type'),'dxa');mar.append(x)
                pr.append(mar)
                if ri==0:x=OxmlElement('w:shd');x.set(qn('w:fill'),'EEEEEE');pr.append(x)
                for p in cell.paragraphs:
                    p.style='Normal'
                    p.paragraph_format.line_spacing=1;p.paragraph_format.space_before=Pt(0);p.paragraph_format.space_after=Pt(0);p.paragraph_format.keep_with_next=False
                    for r in p.runs:r.font.name='Times New Roman';r.font.size=Pt(9.5);r.bold=True if ri==0 else r.bold
    if issup:
        # Place each supplementary figure with its legend in its own landscape page.
        for p in list(d.paragraphs):
            if p.text.startswith('Figure S'):
                p.paragraph_format.page_break_before=True;p.paragraph_format.keep_with_next=True
                fig='FigureS1' if p.text.startswith('Figure S1') else 'FigureS2'
                np=d.add_paragraph();p._p.addnext(np._p)
                run=np.add_run();run.add_picture(str(O/'figures'/(fig+'.png')),width=Inches(8.9 if fig=='FigureS1' else 5.5))
                np.alignment=WD_ALIGN_PARAGRAPH.CENTER
                np.paragraph_format.line_spacing=1;np.paragraph_format.space_after=Pt(0)
    seen=set()
    for sec in d.sections:
        for footer in [sec.footer,sec.first_page_footer,sec.even_page_footer]:
            if footer.part.partname in seen:continue
            seen.add(footer.part.partname)
            for x in list(footer._element):footer._element.remove(x)
            p=footer.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            suppress=OxmlElement('w:suppressLineNumbers');p._p.get_or_add_pPr().append(suppress)
            rr=p.add_run();begin=OxmlElement('w:fldChar');begin.set(qn('w:fldCharType'),'begin');rr._r.append(begin)
            rr=p.add_run();instr=OxmlElement('w:instrText');instr.text=' PAGE ';rr._r.append(instr)
            rr=p.add_run();end=OxmlElement('w:fldChar');end.set(qn('w:fldCharType'),'end');rr._r.append(end)
    # Normalize section-property element order required by WordprocessingML.
    order=['headerReference','footerReference','footnotePr','endnotePr','type','pgSz','pgMar','paperSrc','pgBorders','lnNumType','pgNumType','cols','formProt','vAlign','noEndnote','titlePg','textDirection','bidi','rtlGutter','docGrid','printerSettings']
    for sec in d.sections:
        elements=list(sec._sectPr)
        for el in elements:sec._sectPr.remove(el)
        for el in sorted(elements,key=lambda x: order.index(x.tag.split('}')[-1]) if x.tag.split('}')[-1] in order else 99):sec._sectPr.append(el)
    # Remove generator-specific personal properties and automatic hyperlink styling in titles.
    d.core_properties.last_modified_by='';d.save(out)
    print(name,'tables',len(d.tables),'sections',len(d.sections))
for src,dst in [('instruments_verified.csv','Supplementary_Data_1_Instruments.csv'),('MR_primary_canonical.csv','Supplementary_Data_2_MR.csv'),('coloc_canonical_v2.csv','Supplementary_Data_3_Colocalization.csv')]:shutil.copy2(P/src,O/dst)
